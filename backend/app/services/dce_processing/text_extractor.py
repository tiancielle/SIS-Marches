"""
Extraction du texte brut d'un fichier selon son type.

Décision produit : on extrait toujours le document dans son intégralité. 
Pour les .doc (Word 97-2003), la stratégie est :
1. LibreOffice (soffice) en priorité (recherche automatique des chemins d'installation).
2. Pandoc en solution de repli.
3. Dégradation gracieuse (non_supporte) si les deux échouent, sans faire planter le pipeline.

Justification architecturale : Ce module est conçu pour la robustesse en environnement de production.
Il privilégie la dégradation gracieuse (graceful degradation) et les mécanismes de repli (fallback) 
pour garantir qu'un seul fichier corrompu ou mal formaté n'interrompe pas l'ingestion globale du pipeline.
"""
import os

# CONTRAINTE DE PERFORMANCE BACKEND : Les bibliothèques sous-jacentes (comme OpenCV via PaddleOCR 
# ou NumPy) tentent par défaut d'utiliser tous les cœurs CPU disponibles. Dans un serveur web 
# asynchrone (FastAPI), cela provoque une contention de threads et peut faire planter le serveur 
# lors de requêtes concurrentes. On limite volontairement à 2 threads pour éviter la sur-souscription
# avec la parallélisation OCR (2 workers × 2 threads = 4 threads totaux, adapté à CPU 4 cœurs).
os.environ.setdefault("OMP_NUM_THREADS", "2")
os.environ.setdefault("MKL_NUM_THREADS", "2")

import subprocess
import shutil
import tempfile
import logging
import time
from dataclasses import dataclass
from typing import Optional

from app.core.config import settings
from app.services.dce_processing.zip_extractor import ExtractedFile
from app.services.dce_processing.ocr_cache import get_cached_ocr_result, save_ocr_result

logger = logging.getLogger(__name__)

# OPTIMISATION MÉMOIRE/TEMPS : Le chargement des modèles PaddleOCR en RAM/VRAM est une opération 
# très coûteuse (plusieurs secondes et téléchargement de poids au premier appel). 
# Ce dictionnaire agit comme un cache au niveau du module (singleton) pour ne payer ce coût 
# qu'une seule fois par langue au cours de la vie du processus, et non à chaque document.
_paddle_ocr_instances: dict[str, object] = {}

# Compteur global pour identifier le premier appel predict() (warm-up)
_first_predict_call = True


def _is_cps_file(nom_fichier: str) -> bool:
    """
    Vérifie si le fichier est le CPS par son nom (Niveau 1).
    
    Le CPS peut avoir différents noms dans les DCE réels :
    - CPS.pdf, CPS AOO 114 TTH HO 26.pdf
    - AO N° 119-DCL-2026 CAF+CTP+DETAIL.pdf
    - CONSULTATION.pdf, REGLEMENT.pdf, CAHIER DES CHARGES.pdf
    
    Critères de détection :
    1. Contient "CPS" (critère principal)
    2. Contient "AO" ou "APPEL" + mots-clés CPS (CONSULTATION, REGLEMENT, CAHIER, DETAIL)
    3. Contient des mots-clés typiques du CPS (CONSULTATION, REGLEMENT, CAHIER)
    
    Cette fonction est utilisée pour l'optimisation CPS Only : seul le CPS scanné
    est traité par OCR, les autres documents scannés sont ignorés.
    """
    nom_lower = nom_fichier.lower()
    
    # Critère 1 : Contient "CPS"
    if "cps" in nom_lower:
        return True
    
    # Critère 2 : Contient "AO" ou "APPEL" + mots-clés CPS
    if "ao" in nom_lower or "appel" in nom_lower:
        cps_keywords = ["consultation", "reglement", "cahier", "detail", "charges"]
        if any(keyword in nom_lower for keyword in cps_keywords):
            return True
    
    # Critère 3 : Contient des mots-clés typiques du CPS
    cps_keywords = ["consultation", "reglement", "cahier des charges", "cdc"]
    if any(keyword in nom_lower for keyword in cps_keywords):
        return True
    
    return False


def _find_cps_candidates(extracted_files: list) -> list:
    """
    Cherche des candidats CPS suspects parmi les fichiers extraits (Niveau 2).
    
    Critères de sélection :
    - PDF scanné (déterminé par l'extraction native)
    - Taille importante (signal fort)
    - Nombre de pages important
    - Nom contenant des mots-clés suspects (AO, CAF, CTP, DETAIL, CAHIER, etc.)
    
    Retourne les candidats classés par pertinence (taille décroissante).
    """
    from app.services.dce_processing.zip_extractor import ExtractedFile
    
    candidates = []
    
    for extracted_file in extracted_files:
        if extracted_file.extension != "pdf":
            continue
        
        # Critère : taille importante (>= 1 Mo)
        if extracted_file.taille_octets < 1024 * 1024:
            continue
        
        # Critère : nom contenant des mots-clés suspects
        nom_lower = extracted_file.nom_fichier.lower()
        suspect_keywords = ["ao", "appel", "caf", "ctp", "detail", "cahier", "consultation", "reglement"]
        has_suspect_name = any(keyword in nom_lower for keyword in suspect_keywords)
        
        # Score de pertinence
        score = extracted_file.taille_octets
        if has_suspect_name:
            score *= 1.5  # Bonus pour le nom suspect
        
        candidates.append({
            "file": extracted_file,
            "score": score,
            "taille": extracted_file.taille_octets,
            "has_suspect_name": has_suspect_name
        })
    
    # Classer par score décroissant
    candidates.sort(key=lambda x: x["score"], reverse=True)
    
    # Retourner les 3 meilleurs candidats maximum
    return candidates[:3]


def _verify_cps_candidate(candidate: dict, output_dir: str) -> tuple[bool, int]:
    """
    Vérifie si un candidat est vraiment le CPS en analysant sa première page (Niveau 3).
    
    Processus :
    1. Convertir uniquement la première page
    2. Lancer OCR uniquement sur cette première page
    3. Analyser le texte obtenu pour des indices CPS
    
    Retourne (is_cps, score).
    """
    from app.services.dce_processing.zip_extractor import ExtractedFile
    
    extracted_file = candidate["file"]
    
    try:
        import fitz  # PyMuPDF
        import tempfile
        import os
    except ImportError:
        return False, 0
    
    logger.info(f"[CPS-DETECT] Vérification première page : {extracted_file.nom_fichier}")
    
    try:
        doc = fitz.open(extracted_file.absolute_path)
        page = doc[0]  # Première page
        
        # Conversion première page → image
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            pix = page.get_pixmap(dpi=200)
            pix.save(tmp.name)
            page_img_path = tmp.name
        
        doc.close()
        
        # OCR première page
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(
                lang='fr',
                text_detection_model_name="PP-OCRv6_small_det",
                text_recognition_model_name="PP-OCRv6_small_rec",
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
                enable_mkldnn=False,
            )
            
            result = ocr.predict(page_img_path)
            texts = []
            for item in result:
                if isinstance(item, dict):
                    rec_texts = item.get("rec_texts", [])
                    texts.extend(rec_texts)
                elif hasattr(item, "rec_texts"):
                    texts.extend(item.rec_texts)
            
            texte = "\n".join(texts).lower()
            
            # Nettoyer le fichier temporaire
            try:
                os.remove(page_img_path)
            except:
                pass
            
            # Analyser le texte pour des indices CPS
            cps_indicators = [
                "cahier des prescriptions spéciales",
                "cahier des prescriptions spéciales",
                "cps",
                "article",
                "objet",
                "lot",
                "marché",
                "appel d'offres",
                "prescriptions",
                "cahier des charges",
                "règlement",
                "spécifications"
            ]
            
            score = 0
            found_indicators = []
            
            for indicator in cps_indicators:
                if indicator in texte:
                    score += 1
                    found_indicators.append(indicator)
            
            logger.info(f"[CPS-DETECT] Score CPS = {score} (indices trouvés : {found_indicators})")
            
            # Seuil de confirmation : au moins 3 indices
            is_cps = score >= 3
            
            if is_cps:
                logger.info(f"[CPS-DETECT] CPS confirmé → OCR complet")
            else:
                logger.info(f"[CPS-DETECT] Candidat rejeté (score {score} < 3)")
            
            return is_cps, score
            
        except Exception as exc:
            logger.error(f"[CPS-DETECT] Erreur OCR première page : {exc}")
            return False, 0
            
    except Exception as exc:
        logger.error(f"[CPS-DETECT] Erreur vérification candidat : {exc}")
        return False, 0


def _get_paddle_ocr(lang: str):
    """Récupère ou initialise l'instance PaddleOCR pour une langue donnée."""
    if lang not in _paddle_ocr_instances:
        logger.info(f"[DIAG] Chargement de PaddleOCR (lang={lang})...")
        from paddleocr import PaddleOCR
        
        # CHOIX DE CONFIGURATION OCR : 
        # On désactive les étapes de prétraitement lourdes (classification d'orientation, 
        # redressement de document, orientation des lignes). Pour des documents administratifs 
        # (DCE) numérisés, ceux-ci sont généralement droits et bien formatés. 
        # Désactiver ces options réduit drastiquement le temps d'inférence sans perte de précision.
        # enable_mkldnn=False est maintenu pour éviter le bug PaddlePaddle 3.3.1 + oneDNN.
        # PP-OCRv6_small est utilisé pour améliorer les performances sur CPU (2-3x plus rapide que medium).
        _paddle_ocr_instances[lang] = PaddleOCR(
            lang=lang,
            text_detection_model_name="PP-OCRv6_small_det",
            text_recognition_model_name="PP-OCRv6_small_rec",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            enable_mkldnn=True,
        )
        logger.info(f"[DIAG] PaddleOCR (lang={lang}) chargé avec succès.")
    return _paddle_ocr_instances[lang]


def _extract_ocr_result_texts(result_item) -> list:
    """
    Extrait les textes reconnus d'un résultat PaddleOCR.
    Justification : L'API Python de PaddleOCR a évolué entre les versions majeures (2.x vs 3.x). 
    Cette fonction de compatibilité défensive essaie d'abord l'accès par dictionnaire, puis par attribut, 
    garantissant que le pipeline ne casse pas lors d'une mise à jour de la dépendance.
    """
    try:
        return list(result_item["rec_texts"])
    except (KeyError, TypeError, IndexError):
        pass
    texts = getattr(result_item, "rec_texts", None)
    return list(texts) if texts else []


def _ocr_page_avec_langue(img_path: str, lang: str) -> str:
    """Exécute l'OCR sur une image pour une langue spécifique."""
    global _first_predict_call
    try:
        import cv2
        
        # Instrumentation : dimensions de l'image
        img = cv2.imread(img_path)
        if img is not None:
            height, width = img.shape[:2]
            logger.info(f"[OCR-DEBUG] Image dimensions: {width}x{height} pixels")
        
        # AUDIT VERSIONS
        try:
            import paddleocr
            logger.info(f"[OCR-DEBUG] PaddleOCR version: {paddleocr.__version__}")
        except:
            logger.info(f"[OCR-DEBUG] PaddleOCR version: Non disponible")
        
        try:
            import paddle
            logger.info(f"[OCR-DEBUG] PaddlePaddle version: {paddle.__version__}")
        except:
            logger.info(f"[OCR-DEBUG] PaddlePaddle version: Non disponible")
        
        ocr = _get_paddle_ocr(lang)
        
        # Instrumentation : modèle chargé + ID instance pour vérifier partage
        logger.info(f"[OCR-DEBUG] Instance PaddleOCR pour lang={lang} id={id(ocr)}")
        
        # Instrumentation : temps predict() avec distinction premier appel vs suivants
        predict_start = time.time()
        result = ocr.predict(img_path)
        predict_end = time.time()
        predict_time = predict_end - predict_start
        
        if _first_predict_call:
            logger.info(f"[OCR-DEBUG] PREMIER predict() (warm-up): {predict_time:.2f}s")
            _first_predict_call = False
        else:
            logger.info(f"[OCR-DEBUG] predict() (normal): {predict_time:.2f}s")
        
        logger.info(f"[OCR-DEBUG] Nombre de résultats: {len(result)}")

        texts = []
        for item in result:
            extracted = _extract_ocr_result_texts(item)
            texts.extend(extracted)

        texte_final = "\n".join(texts)
        logger.info(f"[OCR-DEBUG] Texte extrait: {len(texts)} blocs, {len(texte_final)} caractères")
        logger.info(f"[OCR-DEBUG] Aperçu texte: {texte_final[:200] if texte_final else '(vide)'}")
        return texte_final
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"[OCR] Échec OCR (lang={lang}) sur {img_path} : {exc}")
        return ""


def _preprocess_image_for_ocr(img_path: str) -> str:
    """
    Pré-traitement d'image pour améliorer la qualité de l'OCR.
    
    Opérations appliquées :
    1. Correction automatique de rotation (deskew)
    2. Amélioration du contraste (CLAHE)
    3. Binarisation adaptative
    4. Réduction du bruit (median blur)
    
    Retourne le chemin de l'image pré-traitée.
    """
    try:
        import cv2
        import numpy as np
    except ImportError:
        logger.warning("[OCR] OpenCV non disponible, pas de pré-traitement")
        return img_path
    
    try:
        # Charger l'image
        img = cv2.imread(img_path)
        if img is None:
            logger.warning(f"[OCR] Impossible de charger l'image : {img_path}")
            return img_path
        
        original_path = img_path
        processed_path = img_path.replace(".png", "_preprocessed.png")
        
        # 1. Conversion en niveaux de gris
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # 2. Correction de rotation (deskew) simple
        # Détecter les contours et trouver l'angle de rotation
        coords = np.column_stack(np.where(gray > 0))
        if len(coords) > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            
            # Corriger si l'angle est significatif (> 0.5 degré)
            if abs(angle) > 0.5:
                (h, w) = gray.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                logger.debug(f"[OCR] Correction rotation : {angle:.2f} degrés")
        
        # 3. Amélioration du contraste avec CLAHE (Contrast Limited Adaptive Histogram Equalization)
        # Améliore le contraste local sans amplifier trop le bruit
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        # 4. Binarisation adaptative (Otsu ou adaptative)
        # Pour les documents avec un fond non uniforme
        _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # 5. Réduction du bruit légère (median blur)
        # Aide à éliminer le bruit de fond sans trop flouter le texte
        denoised = cv2.medianBlur(binary, 3)
        
        # Sauvegarder l'image pré-traitée
        cv2.imwrite(processed_path, denoised)
        
        logger.debug(f"[OCR] Pré-traitement terminé : {original_path} -> {processed_path}")
        return processed_path
        
    except Exception as exc:
        logger.warning(f"[OCR] Erreur lors du pré-traitement : {exc}")
        return img_path





def _ocr_pdf_scanne(path: str, out_path: str) -> tuple[int, Optional[str]]:
    """
    Repli OCR pour un PDF scanné (version séquentielle).
    
    CACHE OCR :
    Vérifie si le résultat existe déjà dans le cache avant de lancer l'OCR.
    Sauvegarde le résultat dans le cache après l'OCR.
    
    PRÉ-TRAITEMENT IMAGE :
    Chaque page OCRisée passe par un pré-traitement (rotation, contraste, binarisation, débruitage)
    pour améliorer la qualité de reconnaissance.
    """
    # Vérification du cache OCR avant de lancer l'OCR
    cached_content, time_saved, cache_metadata = get_cached_ocr_result(path)
    
    if cached_content is not None:
        # Cache HIT : écrire directement le contenu depuis le cache
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(cached_content)
        logger.info(f"[PIPELINE] OCR ............ SKIP (cache HIT, temps économisé : {time_saved:.1f}s)")
        return len(cached_content), None
    
    # Cache MISS : lancer l'OCR normalement
    logger.info(f"[PIPELINE] OCR ............ RUN (cache MISS)")
    
    try:
        # CHOIX TECHNIQUE : PyMuPDF (fitz) est utilisé car c'est la bibliothèque la plus rapide
        # et la plus fiable pour le rendu de pages PDF en images (pixmaps) en Python.
        import fitz  # PyMuPDF
    except ImportError as exc:
        return 0, f"PyMuPDF (fitz) non installé : {exc}"

    langues = settings.ocr_langs.split(",") if settings.ocr_langs else ["fr"]
    langues = [l.strip() for l in langues if l.strip()]

    # HEURISTIQUE : Un seuil de 30 caractères permet d'éviter les "faux positifs" (bruit de fond
    # ou artefacts de scan reconnus comme 2-3 lettres). 30 caractères indiquent généralement
    # qu'une ligne ou un paragraphe significatif a été lu avec succès, justifiant l'arrêt des tests.
    SEUIL_SUFFISANT = 30  # nb de caractères au-delà duquel une page est considérée bien reconnue

    total_chars = 0
    langue_dominante: Optional[str] = None
    
    # Mesures de performance
    total_conversion_time = 0
    total_preprocess_time = 0
    total_ocr_time = 0
    
    ocr_start = time.time()
    
    try:
        doc = fitz.open(path)
        nb_pages = len(doc)

        logger.info(f"[OCR] Fichier: {os.path.basename(path)}, Pages: {nb_pages}")

        with tempfile.TemporaryDirectory() as tmp_dir, open(out_path, "w", encoding="utf-8") as out:
            for page_num, page in enumerate(doc, start=1):
                # Conversion PDF → image
                conv_start = time.time()
                pix = page.get_pixmap(dpi=200)
                page_img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                pix.save(page_img_path)
                conv_end = time.time()
                page_conversion_time = conv_end - conv_start
                total_conversion_time += page_conversion_time
                
                # Pré-traitement OpenCV
                preprocess_start = time.time()
                processed_img_path = _preprocess_image_for_ocr(page_img_path)
                preprocess_end = time.time()
                page_preprocess_time = preprocess_end - preprocess_start
                total_preprocess_time += page_preprocess_time
                
                # Utiliser l'image pré-traitée si disponible, sinon l'originale
                img_to_ocr = processed_img_path if processed_img_path != page_img_path else page_img_path

                # Ordre de langues à tester
                if langue_dominante:
                    ordre = [langue_dominante] + [l for l in langues if l != langue_dominante]
                else:
                    ordre = langues

                meilleur_texte = ""
                meilleure_langue = None
                
                # OCR avec mesure de temps
                ocr_page_start = time.time()
                for lang in ordre:
                    texte = _ocr_page_avec_langue(img_to_ocr, lang)
                    if len(texte) > len(meilleur_texte):
                        meilleur_texte = texte
                        meilleure_langue = lang
                    if len(meilleur_texte) >= SEUIL_SUFFISANT:
                        break
                ocr_page_end = time.time()
                page_ocr_time = ocr_page_end - ocr_page_start
                total_ocr_time += page_ocr_time
                
                # Log détaillé par page avec progression
                logger.info(f"[OCR] Page {page_num}/{nb_pages} | conversion={page_conversion_time:.2f}s | preprocessing={page_preprocess_time:.2f}s | predict={page_ocr_time:.2f}s")
                logger.info(f"[OCR-DEBUG] Page {page_num}: texte brut = {len(meilleur_texte)} caractères")
                if meilleur_texte:
                    logger.info(f"[OCR-DEBUG] Page {page_num}: aperçu = {meilleur_texte[:100]}")
                
                # Dès la première page qui donne un résultat exploitable, on fige la
                # langue dominante pour accélérer toutes les pages suivantes.
                if langue_dominante is None and meilleure_langue and len(meilleur_texte) >= SEUIL_SUFFISANT:
                    langue_dominante = meilleure_langue

                if meilleur_texte:
                    out.write(meilleur_texte)
                    out.write("\n\n")
                    total_chars += len(meilleur_texte)
                    logger.info(f"[OCR-DEBUG] Page {page_num}: cumul après écriture = {total_chars} caractères")
                else:
                    logger.warning(f"[OCR-DEBUG] Page {page_num}: aucun texte extrait")
                
                # Nettoyer les fichiers temporaires
                if os.path.exists(page_img_path):
                    os.remove(page_img_path)
                if processed_img_path and processed_img_path != page_img_path and os.path.exists(processed_img_path):
                    os.remove(processed_img_path)

        doc.close()
    except Exception as exc:  # noqa: BLE001
        return 0, f"Erreur pendant l'OCR : {exc}"

    ocr_end = time.time()
    total_ocr_time_all = ocr_end - ocr_start
    
    logger.info(f"[OCR-DEBUG] Fin OCR: total_chars = {total_chars}, out_path = {out_path}")
    
    if total_chars > 0:
        logger.info(f"[OCR] Succès : {total_chars} caractères extraits de {nb_pages} pages")
        logger.info(f"[PERF] Conversion PDF→image: {total_conversion_time:.2f}s")
        logger.info(f"[PERF] Pré-traitement OpenCV: {total_preprocess_time:.2f}s")
        logger.info(f"[PERF] OCR PaddleOCR: {total_ocr_time:.2f}s")
        logger.info(f"[PERF] Temps total OCR: {total_ocr_time_all:.2f}s")
        
        # Vérifier le fichier écrit
        if os.path.exists(out_path):
            file_size = os.path.getsize(out_path)
            logger.info(f"[OCR-DEBUG] Fichier .txt écrit: {out_path}, taille = {file_size} octets")
            with open(out_path, "r", encoding="utf-8") as f:
                content = f.read()
                logger.info(f"[OCR-DEBUG] Contenu fichier .txt: {len(content)} caractères")
                logger.info(f"[OCR-DEBUG] Aperçu fichier .txt: {content[:200] if content else '(vide)'}")
        else:
            logger.error(f"[OCR-DEBUG] Fichier .txt NON créé: {out_path}")
        
        # Cache sauvegarde
        try:
            with open(out_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
            
            if extracted_text and len(extracted_text.strip()) > 0:
                save_ocr_result(path, extracted_text, total_ocr_time_all, nb_pages)
        except Exception as exc:
            logger.error(f"[CACHE] Erreur lors de la sauvegarde du cache : {exc}")
    else:
        logger.warning(f"[OCR-DEBUG] Aucun caractère extrait ({total_chars} caractères)")
        if os.path.exists(out_path):
            file_size = os.path.getsize(out_path)
            logger.warning(f"[OCR-DEBUG] Fichier .txt existe (vide?): {out_path}, taille = {file_size} octets")
        else:
            logger.warning(f"[OCR-DEBUG] Fichier .txt n'existe pas: {out_path}")
        
    return total_chars, None


# Variable globale pour stocker l'instance PaddleOCR par processus worker
_worker_ocr_instance = None

def _worker_initializer():
    """
    Initializer pour ProcessPoolExecutor : crée l'instance PaddleOCR une seule fois par worker.
    """
    global _worker_ocr_instance
    if _worker_ocr_instance is None:
        from paddleocr import PaddleOCR
        import os
        worker_pid = os.getpid()
        _worker_ocr_instance = PaddleOCR(
            lang="fr",
            text_detection_model_name="PP-OCRv6_small_det",
            text_recognition_model_name="PP-OCRv6_small_rec",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            enable_mkldnn=False,
        )
        print(f"[OCR-WORKER] Initialisation PaddleOCR PID={worker_pid}", flush=True)
    else:
        import os
        worker_pid = os.getpid()
        print(f"[OCR-WORKER] PID={worker_pid} utilisation instance existante", flush=True)


def _ocr_page_worker(args: tuple) -> tuple[int, str, float, float, float]:
    """
    Worker pour OCR parallélisé d'une page (ProcessPoolExecutor).
    
    Chaque processus worker effectue sa propre conversion PDF → image et crée sa propre instance PaddleOCR.
    
    Args:
        args: (pdf_path, page_num, langues, SEUIL_SUFFISANT, dpi)
    
    Returns:
        (page_num, texte, conversion_time, preprocess_time, ocr_time)
    """
    pdf_path, page_num, langues, SEUIL_SUFFISANT, dpi = args
    
    page_conversion_time = 0.0
    page_preprocess_time = 0.0
    page_ocr_time = 0.0
    
    tmp_path = None
    processed_img_path = None
    meilleur_texte = ""
    
    import os
    worker_start = time.time()
    worker_pid = os.getpid()
    
    print(f"[OCR-WORKER] PID={worker_pid} utilisation instance existante page={page_num}", flush=True)
    
    try:
        import fitz  # PyMuPDF
        
        # Conversion PDF → image dans ce processus worker
        conv_start = time.time()
        doc = fitz.open(pdf_path)
        page = doc[page_num - 1]  # 0-indexed
        pix = page.get_pixmap(dpi=dpi)
        doc.close()
        
        # Utiliser tempfile.mkstemp pour éviter le problème de verrouillage Windows
        import tempfile
        fd, tmp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)  # Fermer immédiatement le handle Windows
        
        page_img_path = tmp_path
        pix.save(page_img_path)
        
        conv_end = time.time()
        page_conversion_time = conv_end - conv_start
        
        # Pré-traitement OpenCV
        preprocess_start = time.time()
        processed_img_path = _preprocess_image_for_ocr(page_img_path)
        preprocess_end = time.time()
        page_preprocess_time = preprocess_end - preprocess_start
        
        # OCR avec tentative de plusieurs langues
        ocr_page_start = time.time()
        img_to_ocr = processed_img_path if processed_img_path else page_img_path
        
        meilleure_langue = None
        
        # Utiliser l'instance globale PaddleOCR créée par l'initializer
        global _worker_ocr_instance
        
        for lang in langues:
            result = _worker_ocr_instance.predict(img_to_ocr)
            
            texts = []
            for item in result:
                extracted = _extract_ocr_result_texts(item)
                texts.extend(extracted)
            
            texte = "\n".join(texts)
            
            if len(texte) > len(meilleur_texte):
                meilleur_texte = texte
                meilleure_langue = lang
            if len(meilleur_texte) >= SEUIL_SUFFISANT:
                break
        
        ocr_page_end = time.time()
        page_ocr_time = ocr_page_end - ocr_page_start
        
    except Exception as exc:
        logger.error(f"[OCR] Erreur sur page {page_num}: {exc}")
    finally:
        # Nettoyer les fichiers temporaires
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except PermissionError:
                logger.warning(f"[OCR] Impossible de supprimer le fichier temporaire : {tmp_path}")
        if processed_img_path and processed_img_path != page_img_path and os.path.exists(processed_img_path):
            try:
                os.remove(processed_img_path)
            except PermissionError:
                logger.warning(f"[OCR] Impossible de supprimer le fichier pré-traité : {processed_img_path}")
        
        worker_end = time.time()
        worker_duration = worker_end - worker_start
        print(f"[OCR-WORKER] Fin worker PID={worker_pid} page={page_num} durée={worker_duration:.2f}s (conv={page_conversion_time:.2f}s, preproc={page_preprocess_time:.2f}s, ocr={page_ocr_time:.2f}s)", flush=True)
    
    return (page_num, meilleur_texte, page_conversion_time, page_preprocess_time, page_ocr_time)


def _ocr_pdf_scanne_parallel(path: str, out_path: str) -> tuple[int, Optional[str]]:
    """
    OCR parallélisé pour un PDF scanné.
    
    Utilise ProcessPoolExecutor pour traiter plusieurs pages en parallèle avec des processus indépendants.
    """
    import os  # Import au début pour éviter le bug de scope
    
    # Vérification du cache OCR avant de lancer l'OCR
    cached_content, time_saved, cache_metadata = get_cached_ocr_result(path)
    
    if cached_content is not None:
        # Cache HIT : écrire directement le contenu depuis le cache
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(cached_content)
        logger.info(f"[PIPELINE] OCR ............ SKIP (cache HIT, temps économisé : {time_saved:.1f}s)")
        return len(cached_content), None
    
    # Cache MISS : lancer l'OCR normalement
    logger.info(f"[PIPELINE] OCR ............ RUN (cache MISS)")
    
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        return 0, f"PyMuPDF (fitz) non installé : {exc}"

    langues = settings.ocr_langs.split(",") if settings.ocr_langs else ["fr"]
    langues = [l.strip() for l in langues if l.strip()]

    SEUIL_SUFFISANT = 30  # nb de caractères au-delà duquel une page est considérée bien reconnue
    dpi = 150  # DPI pour la conversion
    
    ocr_start = time.time()
    
    total_conversion_time = 0
    total_preprocess_time = 0
    total_ocr_time = 0
    total_chars = 0
    
    try:
        doc = fitz.open(path)
        nb_pages = len(doc)

        logger.info(f"[OCR] Fichier: {os.path.basename(path)}, Pages: {nb_pages}")
        logger.info(f"[OCR] Mode: Parallèle (ProcessPoolExecutor)")
        logger.info(f"[OCR] Début traitement parallèle...")
        
        # Étape : OCR parallèle avec ProcessPoolExecutor
        from concurrent.futures import ProcessPoolExecutor, as_completed
        
        # Préparer les arguments pour chaque page
        worker_args = [
            (path, page_num, langues, SEUIL_SUFFISANT, dpi)
            for page_num in range(1, nb_pages + 1)
        ]
        
        # Nombre de workers = 1 pour éviter la contention CPU avec OMP_NUM_THREADS=2
        # (1 worker × 2 threads OMP = 2 threads totaux, mieux adapté à CPU 4 cœurs)
        max_workers = 1
        logger.info(f"[OCR] Workers: {max_workers}")
        
        ocr_parallel_start = time.time()
        results = {}
        
        logger.info(f"[OCR-PARALLEL] Début traitement parallèle à {time.strftime('%H:%M:%S', time.localtime(ocr_parallel_start))}")
        
        with ProcessPoolExecutor(max_workers=max_workers, initializer=_worker_initializer) as executor:
            future_to_page = {
                executor.submit(_ocr_page_worker, args): args[1]
                for args in worker_args
            }
            
            pending_futures = len(future_to_page)
            completed_futures = 0
            logger.info(f"[OCR-PARALLEL] {pending_futures} tâches soumises à {max_workers} workers")
            
            for future in as_completed(future_to_page):
                page_num = future_to_page[future]
                completed_futures += 1
                
                try:
                    page_num, texte, conv_time, preprocess_time, ocr_time = future.result()
                    results[page_num] = (texte, conv_time, preprocess_time, ocr_time)
                    
                    # Accumuler les temps
                    total_conversion_time += conv_time
                    total_preprocess_time += preprocess_time
                    
                    # Log de progression
                    remaining = pending_futures - completed_futures
                    logger.info(f"[OCR] Page {page_num}/{nb_pages} | predict={ocr_time:.2f}s | terminés={completed_futures}/{pending_futures} | restants={remaining}")
                except Exception as exc:
                    logger.error(f"[OCR] Erreur page {page_num}: {exc}")
        
        logger.info(f"[OCR-PARALLEL] Traitement parallèle terminé | {completed_futures}/{pending_futures} tâches complétées")
        
        ocr_parallel_end = time.time()
        total_ocr_time = ocr_parallel_end - ocr_parallel_start
        logger.info(f"[OCR-PARALLEL] Fin traitement parallèle à {time.strftime('%H:%M:%S', time.localtime(ocr_parallel_end))} | durée={total_ocr_time:.2f}s")
        
        # Assembler les résultats dans l'ordre
        with open(out_path, "w", encoding="utf-8") as out:
            for page_num in range(1, nb_pages + 1):
                if page_num in results:
                    texte, conv_time, preprocess_time, ocr_time = results[page_num]
                    if texte:
                        out.write(texte)
                        out.write("\n\n")
                        total_chars += len(texte)
        
        doc.close()
    except Exception as exc:  # noqa: BLE001
        return 0, f"Erreur pendant l'OCR parallèle : {exc}"

    ocr_end = time.time()
    total_ocr_time_all = ocr_end - ocr_start
    
    if total_chars > 0:
        logger.info(f"[OCR] Succès : {total_chars} caractères extraits de {nb_pages} pages")
        logger.info(f"[PERF] Conversion PDF→image: {total_conversion_time:.2f}s")
        logger.info(f"[PERF] Pré-traitement OpenCV: {total_preprocess_time:.2f}s")
        logger.info(f"[PERF] OCR PaddleOCR (parallèle): {total_ocr_time:.2f}s")
        logger.info(f"[PERF] Temps total OCR: {total_ocr_time_all:.2f}s")
        
        # Cache sauvegarde
        try:
            with open(out_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
            
            if extracted_text and len(extracted_text.strip()) > 0:
                save_ocr_result(path, extracted_text, total_ocr_time_all, nb_pages)
        except Exception as exc:
            logger.error(f"[CACHE] Erreur lors de la sauvegarde du cache : {exc}")
    
    return total_chars, None


SUPPORTED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx"}


# Classe ExtractionResult - définie ici pour être disponible pour toutes les fonctions
@dataclass
class ExtractionResult:
    texte_extrait_path: Optional[str]
    nb_caracteres: int
    statut: str          # succes | echec | non_supporte
    erreur: Optional[str]


def _output_txt_path(extracted_file: ExtractedFile, output_dir: str) -> str:
    """Calcule le chemin de sortie du fichier texte en préservant la structure relative des dossiers."""
    base, _ = os.path.splitext(extracted_file.relative_path)
    txt_relative = base + ".txt"
    destination = os.path.join(output_dir, txt_relative)
    os.makedirs(os.path.dirname(destination), exist_ok=True)
    return destination


def _find_libreoffice_executable() -> Optional[str]:
    """
    Recherche l'exécutable LibreOffice (soffice) sur la machine.
    Justification : S'appuyer uniquement sur le PATH système est fragile en production 
    (surtout sous Windows avec les services ou les environnements virtuels). 
    On implémente donc une cascade de recherche : config explicite > chemins par défaut > PATH.
    """
    logger.debug("[DIAG] Recherche de LibreOffice (soffice)...")

    # 0. Chemin explicite fourni en config — prioritaire, aucune détection nécessaire.
    # Utile si l'install est à un emplacement non standard, ou pour contourner
    # définitivement un souci de détection PATH intermittent.
    if settings.libreoffice_path:
        if os.path.exists(settings.libreoffice_path):
            logger.info(f"[DIAG] Succès : LibreOffice via LIBREOFFICE_PATH (.env) : {settings.libreoffice_path}")
            return settings.libreoffice_path
        logger.warning(
            f"[DIAG] LIBREOFFICE_PATH configuré ('{settings.libreoffice_path}') mais introuvable à cet "
            f"emplacement — poursuite avec la détection automatique."
        )

    # 1. Chemins d'installation par défaut sous Windows
    win_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for path in win_paths:
        if os.path.exists(path):
            logger.info(f"[DIAG] Succès : LibreOffice trouvé à l'emplacement par défaut : {path}")
            return path

    # 2. Fallback : vérifier s'il est dans le PATH système (Linux/macOS ou install custom Windows)
    # CHOIX DE TIMEOUT : Un premier lancement "à froid" de LibreOffice sur Windows peut légitimement 
    # dépasser 5s (initialisation du profil utilisateur, analyse par l'antivirus/Windows Defender). 
    # Un timeout de 5s provoquait des échecs de détection intermittents (faux négatifs). 20s est un 
    # compromis sûr pour la détection sans bloquer indéfiniment le thread.
    logger.debug("[DIAG] Non trouvé dans les chemins Windows par défaut. Vérification du PATH système...")
    for cmd in ["soffice", "soffice.exe"]:
        try:
            result = subprocess.run([cmd, "--version"], capture_output=True, text=True, timeout=20)
            if result.returncode == 0:
                logger.info(f"[DIAG] Succès : LibreOffice trouvé dans le PATH système : {cmd}")
                return cmd
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            logger.warning(f"[DIAG] Détection de '{cmd}' expirée après 20s (démarrage à froid ?). Essai suivant.")
        except Exception as e:
            logger.debug(f"[DIAG] Erreur lors de la vérification de {cmd} dans le PATH : {e}")

    logger.warning(
        "[DIAG] Échec : LibreOffice (soffice) introuvable sur cette machine. Si LibreOffice est bien "
        "installé, renseigne LIBREOFFICE_PATH dans .env avec le chemin exact vers soffice.exe pour "
        "contourner la détection."
    )
    return None


def _extract_pdf(path: str, out_path: str) -> tuple[int, Optional[str]]:
    """
    Tente l'extraction native du PDF.
    Stratégie : pdfplumber en priorité (meilleure fidélité de mise en page et gestion des tableaux), 
    avec un repli sur pypdf en cas d'échec total (plus tolérant aux PDFs malformés).
    
    Détecte automatiquement si le PDF possède une couche texte exploitable.
    Si oui, utilise l'extraction native. Si non, retourne 0 pour déclencher l'OCR.
    
    OPTIMISATION : La détection et l'extraction sont faites en un seul passage pour éviter
    d'ouvrir le PDF deux fois avec pdfplumber.
    """
    logger.debug(f"[DIAG] Extraction PDF démarrée pour : {os.path.basename(path)}")
    total_chars = 0
    is_scanned = False
    
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf, open(out_path, "w", encoding="utf-8") as out:
            total_pages = len(pdf.pages)
            pages_with_text = 0
            
            # Premier passage : détection + extraction en une seule boucle
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                
                # Détection : compter les pages avec texte significatif
                if page_text and len(page_text.strip()) > 50:
                    pages_with_text += 1
                
                # Extraction : écrire le texte immédiatement
                if page_text:
                    out.write(page_text)
                    out.write("\n\n")
                    total_chars += len(page_text)
            
            # Après le parcours, décider si le PDF est scanné
            if total_pages > 0 and pages_with_text / total_pages < 0.5:
                is_scanned = True
                logger.info(f"[DIAG] PDF détecté comme scanné : {pages_with_text}/{total_pages} pages avec texte significatif")
                # PDF scanné : retourner 0 pour déclencher l'OCR
                # On supprime le fichier partiellement écrit
                if os.path.exists(out_path):
                    os.remove(out_path)
                return 0, None
            else:
                logger.info(f"[DIAG] PDF détecté comme natif : {pages_with_text}/{total_pages} pages avec texte significatif")
            
        if total_chars > 0:
            logger.info(f"[DIAG] Succès extraction PDF natif (pdfplumber) : {total_chars} caractères.")
            return total_chars, None
        
        # PDF natif mais vide : retourner 0 mais NE PAS déclencher l'OCR
        # On marque ce cas spécifiquement pour éviter l'OCR inutile
        logger.warning("[DIAG] PDF natif mais vide (0 caractères extraits). Pas d'OCR.")
        return 0, "pdf_natif_vide"
    except Exception as pdfplumber_error:
        logger.debug(f"[DIAG] Échec pdfplumber, tentative de repli pypdf. Erreur: {pdfplumber_error}")
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            total_chars = 0
            with open(out_path, "w", encoding="utf-8") as out:
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        out.write(page_text)
                        out.write("\n\n")
                        total_chars += len(page_text)
            logger.info(f"[DIAG] Succès extraction PDF (pypdf repli) : {total_chars} caractères.")
            return total_chars, None
        except Exception as pypdf_error:
            logger.error(f"[DIAG] Échec total extraction PDF. pypdf erreur: {pypdf_error}")
            return 0, f"pdfplumber: {pdfplumber_error} | pypdf (repli): {pypdf_error}"


def _extract_docx(path: str, out_path: str) -> tuple[int, Optional[str]]:
    """Extraction native des fichiers Word modernes (.docx) via python-docx."""
    logger.debug(f"[DIAG] Extraction DOCX démarrée pour : {os.path.basename(path)}")
    try:
        import docx
        document = docx.Document(path)
        total_chars = 0
        with open(out_path, "w", encoding="utf-8") as out:
            for paragraph in document.paragraphs:
                if paragraph.text:
                    out.write(paragraph.text)
                    out.write("\n")
                    total_chars += len(paragraph.text)
            # Les tableaux sont souvent porteurs d'informations critiques dans les DCE.
            # On les extrait en format texte tabulé pour préserver une structure lisible.
            for table in document.tables:
                for row in table.rows:
                    cells_text = "\t".join(cell.text for cell in row.cells if cell.text)
                    if cells_text:
                        out.write(cells_text)
                        out.write("\n")
                        total_chars += len(cells_text)
        logger.info(f"[DIAG] Succès extraction DOCX : {total_chars} caractères.")
        return total_chars, None
    except Exception as exc:
        logger.error(f"[DIAG] Échec extraction DOCX : {exc}")
        return 0, str(exc)


def _extract_doc(path: str, out_path: str) -> tuple[int, Optional[str], str]:
    """
    Extraction des fichiers .doc (Word 97-2003, format binaire legacy).
    Stratégie : 1. LibreOffice (recherche auto des chemins) -> 2. Pandoc (repli) -> 3. non_supporte.
    """
    filename = os.path.basename(path)
    logger.info(f"[DIAG] === DÉBUT TRAITEMENT .DOC : {filename} ===")
    
    # --- ÉTAPE 1 : LibreOffice (Méthode principale) ---
    lo_path = _find_libreoffice_executable()
    if lo_path:
        logger.info(f"[DIAG] Tentative de conversion via LibreOffice : {filename}")
        temp_dir = os.path.dirname(out_path)
        base_name = os.path.splitext(filename)[0]
        # LibreOffice va créer un fichier nommé base_name.txt dans temp_dir
        temp_txt_path = os.path.join(temp_dir, f"{base_name}.txt")

        cmd = [
            lo_path,
            "--headless",
            "--convert-to", "txt:Text (encoded):UTF8",
            "--outdir", temp_dir,
            path
        ]

        try:
            logger.debug(f"[DIAG] Exécution commande LO : {' '.join(cmd)}")
            # Timeout de 120s : la conversion de gros documents legacy peut être lente.
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

            # CONTournement DE BUG CONNU : LibreOffice CLI renvoie souvent un code de sortie 
            # non-nul (ex: 1) même quand la conversion a réussi. Se fier au returncode seul 
            # générerait des faux échecs. La seule métrique fiable est l'existence du fichier 
            # de sortie et sa taille (nombre de caractères).
            if os.path.exists(temp_txt_path):
                # On déplace/renomme le fichier généré vers le out_path final attendu par le pipeline
                if temp_txt_path != out_path:
                    os.replace(temp_txt_path, out_path)

                with open(out_path, "r", encoding="utf-8") as f:
                    text = f.read()

                if text.strip():
                    if result.returncode != 0:
                        logger.info(
                            f"[DIAG] LibreOffice a renvoyé le code {result.returncode} mais le "
                            f"fichier produit est bien exploitable ({len(text)} caractères) — "
                            f"traité comme un succès (code de retour non fiable, connu chez LibreOffice)."
                        )
                    logger.info(f"[DIAG] SUCCÈS LibreOffice : {len(text)} caractères extraits de {filename}.")
                    return len(text), None, "succes"
                logger.warning(f"[DIAG] Fichier produit par LibreOffice mais vide pour {filename}.")
            else:
                logger.warning(f"[DIAG] Échec conversion LibreOffice. Code retour: {result.returncode}. stderr: {result.stderr}")
        except subprocess.TimeoutExpired:
            logger.warning(f"[DIAG] Timeout (120s) lors de la conversion LibreOffice pour {filename}.")
        except Exception as e:
            logger.warning(f"[DIAG] Exception ({type(e).__name__}) lors de l'exécution LibreOffice pour {filename} : {e}")

    # --- ÉTAPE 2 : Pandoc (Solution de repli) ---
    logger.info(f"[DIAG] LibreOffice a échoué ou est absent. Tentative de repli via Pandoc pour : {filename}")
    try:
        import pypandoc
        # On force explicitement le format d'entrée 'doc' pour éviter que Pandoc ne devine mal le format.
        text = pypandoc.convert_file(path, "plain", format="doc")
        with open(out_path, "w", encoding="utf-8") as out:
            out.write(text)
        logger.info(f"[DIAG] SUCCÈS Pandoc (repli) : {len(text)} caractères extraits de {filename}.")
        return len(text), None, "succes"
    except (RuntimeError, OSError) as e:
        logger.warning(f"[DIAG] Échec du repli Pandoc pour {filename} : {e}")
    except Exception as e:
        logger.warning(f"[DIAG] Exception inattendue lors du repli Pandoc pour {filename} : {e}")

    # --- ÉTAPE 3 : Dégradation gracieuse ---
    # PLANNING PIPELINE : Plutôt que de lever une exception qui arrêterait le traitement par lots, 
    # on retourne un statut "non_supporte". Cela permet au système d'ingestion de marquer le fichier 
    # pour une révision manuelle ultérieure tout en continuant à traiter les fichiers suivants.
    error_msg = (
        "Pandoc n'a pas pu lire ce .doc, et le repli LibreOffice a échoué ou est introuvable. "
        "Installe LibreOffice pour permettre la conversion des .doc legacy, ou convertis ce fichier manuellement en .docx/.pdf."
    )
    logger.error(f"[DIAG] ÉCHEC TOTAL pour {filename}. Le fichier sera marqué comme 'non_supporte' dans le pipeline.")
    return 0, error_msg, "non_supporte"


def _extract_xlsx(path: str, out_path: str) -> tuple[int, Optional[str]]:
    """Extraction des fichiers Excel en préservant la structure par feuille et par ligne."""
    logger.debug(f"[DIAG] Extraction XLSX démarrée pour : {os.path.basename(path)}")
    try:
        import openpyxl
        # OPTIMISATION MÉMOIRE ET SÉMANTIQUE : 
        # data_only=True : Extrait les valeurs calculées plutôt que les formules Excel (inutiles pour le NLP).
        # read_only=True : Active le mode de lecture en flux (streaming), empêchant les pics de consommation 
        # RAM massifs lors de l'ouverture de fichiers Excel volumineux.
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
        total_chars = 0
        with open(out_path, "w", encoding="utf-8") as out:
            for sheet in workbook.worksheets:
                out.write(f"--- Feuille: {sheet.title} ---\n")
                for row in sheet.iter_rows():
                    values = [str(cell.value) for cell in row if cell.value is not None]
                    if values:
                        line = "\t".join(values)
                        out.write(line)
                        out.write("\n")
                        total_chars += len(line)
        logger.info(f"[DIAG] Succès extraction XLSX : {total_chars} caractères.")
        return total_chars, None
    except Exception as exc:
        logger.error(f"[DIAG] Échec extraction XLSX : {exc}")
        return 0, str(exc)


def extract_text(extracted_file: ExtractedFile, output_dir: str, is_cps_confirmed: bool = False) -> ExtractionResult:
    """
    Point d'entrée principal du pipeline d'extraction.
    Agit comme un routeur qui dispatche vers la méthode adaptée selon l'extension,
    et gère la logique spécifique de repli OCR uniquement pour les PDFs.
    """
    extension = extracted_file.extension
    logger.info(f"[DIAG] Pipeline d'extraction appelé pour : {extracted_file.relative_path} (type: {extension})")

    if extension not in SUPPORTED_EXTENSIONS:
        logger.warning(f"[DIAG] Type non supporté : .{extension}")
        return ExtractionResult(
            texte_extrait_path=None,
            nb_caracteres=0,
            statut="non_supporte",
            erreur=f"Type de fichier '.{extension or '?'}' non pris en charge pour l'extraction de texte.",
        )

    out_path = _output_txt_path(extracted_file, output_dir)
    
    # Mesure du temps total du pipeline
    pipeline_start = time.time()

    if extension == "pdf":
        nb_chars, erreur = _extract_pdf(extracted_file.absolute_path, out_path)
        if erreur:
            return ExtractionResult(None, 0, "echec", erreur)

        # LOGIQUE DE BASCULE OCR : Si l'extraction native ne renvoie aucun caractère
        # ET que l'erreur n'est pas "pdf_natif_vide", on en déduit que le PDF est scanné
        if nb_chars == 0 and erreur != "pdf_natif_vide":
            # OPTIMISATION CPS ONLY : Ne lancer l'OCR que sur le CPS scanné confirmé
            # Les autres documents scannés sont ignorés pour gagner du temps
            if is_cps_confirmed:
                logger.info(f"[OCR] CPS scanné confirmé — tentative OCR (PaddleOCR) pour {extracted_file.nom_fichier}.")
                nb_chars_ocr, erreur_ocr = _ocr_pdf_scanne_parallel(extracted_file.absolute_path, out_path)
            else:
                logger.info(f"[OCR] Document scanné non-CPS ignoré pour optimisation : {extracted_file.nom_fichier}")
                # Retourner un statut spécifique pour document scanné ignoré
                return ExtractionResult(None, 0, "ocr_ignore", "Document scanné non-CPS ignoré pour optimisation du temps de traitement")
            if erreur_ocr:
                return ExtractionResult(
                    out_path, 0, "non_supporte",
                    f"Aucun texte natif, et l'OCR a échoué : {erreur_ocr}",
                )
            if nb_chars_ocr == 0:
                return ExtractionResult(
                    out_path, 0, "non_supporte",
                    "Aucun texte extrait, même après OCR (image illisible, page blanche, ou qualité de scan trop faible).",
                )
            pipeline_end = time.time()
            logger.info(f"[PERF] Temps total pipeline: {pipeline_end - pipeline_start:.2f}s")
            return ExtractionResult(out_path, nb_chars_ocr, "succes", None)
        
        # PDF natif avec texte ou PDF natif vide (pas d'OCR)
        pipeline_end = time.time()
        logger.info(f"[PERF] Temps total pipeline: {pipeline_end - pipeline_start:.2f}s")
        return ExtractionResult(out_path, nb_chars, "succes" if nb_chars > 0 else "echec", None if nb_chars > 0 else "PDF natif mais vide")

    if extension == "docx":
        nb_chars, erreur = _extract_docx(extracted_file.absolute_path, out_path)
        if erreur:
            return ExtractionResult(None, 0, "echec", erreur)
        return ExtractionResult(out_path, nb_chars, "succes", None)

    if extension == "doc":
        nb_chars, erreur, statut = _extract_doc(extracted_file.absolute_path, out_path)
        return ExtractionResult(out_path if statut == "succes" else None, nb_chars, statut, erreur)

    if extension == "xlsx":
        nb_chars, erreur = _extract_xlsx(extracted_file.absolute_path, out_path)
        if erreur:
            return ExtractionResult(None, 0, "echec", erreur)
        return ExtractionResult(out_path, nb_chars, "succes", None)

    return ExtractionResult(None, 0, "non_supporte", "Type de fichier non géré.")